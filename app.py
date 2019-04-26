#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
import docx
from docx import Document
    
fname = 'file.docx'
document = Document(fname)

with open("list.txt", "r") as list:
    for person in list:
        for p in document.paragraphs:
            if  p.text == 'استاد گرانقدر، جناب آقای تست':
                p.text = ""
                p.add_run('استاد گرانقدر، جناب آقای ' + person.strip()).font.name = 'B Nazanin'
                fname = person + '.docx'
                document.save(fname.strip())
                p.text = 'استاد گرانقدر، جناب آقای تست'
                print(person);
